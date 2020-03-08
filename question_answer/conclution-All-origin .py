import os
import re
from docx import Document




# 获取文件路径
path = "aricar"
files= os.listdir(path)
#print(files)




# 患者
def search_patient(txt):

    a = re.search('.*?' + '\s' + 'eligible' + '\s?' + '.*?\D\.', txt, re.IGNORECASE)
    b = re.search('Eligible' + '\s'+'.*?patients.*?' + '.*?\D\.', txt, re.IGNORECASE)
    # c = re.search(('patients with'  +'.*?stage?'+ '.*?\D\.')+'.*?eligible.*?\D\.', txt,re.I | re.IGNORECASE)
    c = re.search('patients with' + '.*?stage?' + '.*?\D\.', txt, re.I | re.IGNORECASE)
    c1 = re.search('\.(.*?patients with' + '.*?tumours' + '.*?\D\.)', txt, re.I | re.IGNORECASE)
    d = re.search('Patients who had' + '\s' + '.*?\D\.', txt,re.I | re.IGNORECASE)
    if c:
        searchtxt = c
        return searchtxt.group()
    elif c1:
        searchtxt = c1
        return searchtxt.group()
    elif d:
        searchtxt = d
        return searchtxt.group()
    elif b:
        searchtxt = b
        return searchtxt.group()
    elif a:
        searchtxt = a
        return searchtxt.group()
    else:
        return 'No Match'


# 基线特征
def search_characteristics(txt):
    if re.search('.*?'+'\s'+'characteristics'+'\s?'+'.*?\.',txt,re.I |re.IGNORECASE):
        searchtxt = re.search('.*?'+'\s'+'characteristics'+'\s?'+'.*?\.',txt,re.I |re.IGNORECASE)
    else:
        searchtxt =re.search('.*?'+'\s'+'characteristics',txt,re.I |re.IGNORECASE)
    return searchtxt.group()

# 药物
def search_medicine(txt):
    # 打开药物词库
    # h = open('trial_drugbank.txt')
    # res = h.read().splitlines()
    # res = [i.replace('"', '') for i in res]

    # h = open('drugbank_main.txt', encoding='utf-8')
    # res = h.read().splitlines()
    #
    # #print(res)
    # res2 = []
    # for i in res:
    #     newi = re.search('\t(.*)', i)
    #     res2.append(newi.group().replace('\t', '').replace('"', ''))
    h = open('new_drug.txt', encoding='utf-8')
    res = h.read().splitlines()
    medcine = []
    reason = []
    num = 0
    for i in res:
        # a = re.search('.*?' + '\s?' + i + '\s?' + '.*?\.', txt, re.I | re.IGNORECASE) #版本1，废弃
        a = re.search('.*?' + '\s?' + i + '\s?', txt)
        if a:
            #print(i)
            searchtxt = a
            #print(searchtxt.group())
            medcine.append(i)
            reason.append(searchtxt.group())
        #print('\r'+str(num / len(res) * 100) + '%',end='')
        num += 1
    return [medcine,reason]

# 样本量
def search_num(txt):
    searchtxt = ''
    a = re.search('.*?' + '\s' + '\d+\spatients' + '\s?' + '.*?\D\.', txt, re.IGNORECASE)
    if a:
        searchtxt = a
        return searchtxt.group()
    else:
        return 'No Match'

# 线数
def search_phase(txt):
    a = re.search('.*?' + '\s' + 'phase I' + '\s?' + '.*?\D\.', txt,  re.IGNORECASE)
    if a:
        searchtxt = a
        return searchtxt.group()
    else:
        return 'No Match'

# 实验目的
def search_purpose(txt):
    a = re.search('.*?' + '\s' + 'background' + '\s?' + '.*?\..*?\.', txt,  re.I | re.IGNORECASE)
    b = re.search('background' + '\s' + '.*?\..*?\.', txt, re.I | re.IGNORECASE)
    if b:
        searchtxt = b
        return searchtxt.group()
    elif a:
        searchtxt = a
        return searchtxt.group()
    else:
        return 'No Match'

# 研究结果
def search_res(txt):
    a = re.search('.*?' + '\s' + 'findings' + '\s?' + '.*?\..*?\D\.', txt, re.I | re.IGNORECASE)
    b = re.search('\s' + 'results' + '\s?' + '.*?\..*?\D\.', txt, re.I | re.IGNORECASE)
    if b:
        searchtxt = b
        return searchtxt.group()

    elif a:
        searchtxt = a
        return searchtxt.group()
    else:
        return 'No Match'

# 研究结论
def search_con(txt):
    a = re.search('conclusions' + '\s?' + '.*?\D\.', txt, re.I | re.IGNORECASE)
    b = re.search('.*?' + '\s' + 'INTERPRETATION' + '\s?' + '.*?\D\.', txt, re.I | re.IGNORECASE)
    if a:
        searchtxt = a
        return searchtxt.group()

    elif b:
        searchtxt = b
        return searchtxt.group()
    else:
        return 'No Match'

def search_disscussion(txt):
    h = open('discussion.txt', encoding='gbk')
    res = h.read().splitlines()
    #print(res)
    temp = []
    for i in res:
        a = re.search('[)](.*)([(])?', i)
        b = a.group().replace(')', '')
        c = re.sub('（.*?）', '', b)
        temp.append(c)
    keyword = []
    reason = []
    num = 0
    for i in temp:
        a = re.search('.*?' + '\s' + i + '\s?' + '.*?\.', txt, re.I | re.IGNORECASE)
        if a:
            searchtxt = a
            keyword.append(i)
            reason.append(searchtxt.group())
        #print('\r'+str(num / len(res) * 100) + '%',end='')
        num += 1
    return [keyword,reason]


# 遍历文件
for i in files:
    document = Document()
    filename = i
    document.add_heading(i, 0)
    f = open(path+'/'+i,'r',encoding='utf-8')
    txt = f.read()
    f.close()
    #print('\n'+i+'-----------------------------------------')

    #print('搜寻相关药物')
    # medcine_reason = search_medicine(txt)
    #
    # document.add_heading(u'研究相关药物',1)
    # document.add_paragraph(str(medcine_reason[0]))

    patients = search_patient(txt)
    document.add_heading(u'研究患者', 1)
    document.add_paragraph(patients)

    num = search_num(txt)
    document.add_heading(u'样本量', 1)
    document.add_paragraph(num)

    characteristics = search_characteristics(txt)
    document.add_heading(u'基线特征', 1)
    document.add_paragraph(characteristics)

    phase = search_phase(txt)
    document.add_heading(u'试验设计', 1)
    document.add_paragraph(phase)

    purpose = search_purpose(txt)
    document.add_heading(u'研究背景', 1)
    document.add_paragraph(purpose)

    result = search_res(txt)
    document.add_heading(u'研究结果', 1)
    document.add_paragraph(result)

    conclution = search_con(txt)
    document.add_heading(u'研究结论', 1)
    document.add_paragraph(conclution)

    #print('搜寻讨论')
    # discussion =search_disscussion(txt)
    # document.add_heading(u'讨论关键词及讨论内容', 1)
    # document.add_paragraph(str(discussion))


    document.save('result/'+filename+'.docx')








    #print('--------样本量-----------')
    #print(num)
    #print('------ 基线特征-----')
    #print(characteristics)
    #print('-----试验设计------')
    #print(phase)
    #print('------结果-----')
    #print(result)
    #print('------结论-----')
    #print(conclution)
    #print('------讨论-----')
    #print(discussion)

