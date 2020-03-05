import os
import re
from docx import Document
from docx.shared import Inches
from pdf_to_txt import pdf_title
# 获取文件路径
path = "./question_answer/aricar"
files= os.listdir(path)
# print(files)

class Qurey_answer(object):

    def __init__(self,txt):
        self.txt = txt

    # 患者
    def search_patient(self):

        '''a = re.search('.*?' + '\s' + 'eligible' + '\s?' + '.*?\D\.', txt, re.IGNORECASE)
        b = re.search('Eligible' + '\s'+'.*?patients.*?' + '.*?\D\.', txt, re.IGNORECASE)
        # c = re.search(('patients with'  +'.*?stage?'+ '.*?\D\.')+'.*?eligible.*?\D\.', txt,re.I | re.IGNORECASE)
        c = re.search('patients with' + '.*?stage?' + '.*?\D\.', txt, re.I | re.IGNORECASE)
        c1 = re.search('\.(.*?patients with' + '.*?tumours' + '.*?\D\.)', txt, re.I | re.IGNORECASE)
        d = re.search('Patients who had' + '\s' + '.*?\D\.', txt,re.I | re.IGNORECASE)'''# 版本1，弃用

        searcha = re.search('.*?' + '\s' + 'eligible' + '\s?' + '.*?\.', self.txt, re.I | re.IGNORECASE)
        strict1 = re.search('.*?' + '\s' + 'patients who had' + '\s?' + '.*?\.', self.txt, re.I | re.IGNORECASE)
        strict2 = re.search('.*?' + '\s' + 'patients with' + '\s?' + '.*?\.', self.txt, re.I | re.IGNORECASE)

        if searcha:
            searchtxt = searcha
            return searchtxt.group()
        elif strict1 or strict2:
            temp = []
            if strict1:
                temp.append('\n>>>>>>>>>>>>>>>>>>>>>>>>')
                temp.append('\n')
                temp.append(strict1.group())
                temp.append('\n')
                temp.append('>>>>>>>>>>>>>>>>>>>>>>>>')
            if strict2:
                temp.append('\n>>>>>>>>>>>>>>>>>>>>>>>>')
                temp.append('\n')
                temp.append(strict2.group())
                temp.append('\n')
                temp.append('\n>>>>>>>>>>>>>>>>>>>>>>>>')
            return temp
        else:
            return 'No Match'


    # 基线特征
    def search_characteristics(self):
        try:
            if re.search('.*?'+'\s'+'characteristics'+'\s?'+'.*?\.',self.txt,re.I |re.IGNORECASE):
                searchtxt = re.search('.*?'+'\s'+'characteristics'+'\s?'+'.*?\.',self.txt,re.I |re.IGNORECASE)
            else:
                searchtxt =re.search('.*?'+'\s'+'characteristics',self.txt,re.I |re.IGNORECASE)
            return searchtxt.group()
        except Exception as e:
            print(e)
            return []

    # 药物
    def search_medicine(self):
        # 打开药物词库
        # h = open('trial_drugbank.txt')
        # res = h.read().splitlines()
        # res = [i.replace('"', '') for i in res]

        # h = open('drugbank_main.txt', encoding='utf-8')
        # res = h.read().splitlines()
        #
        # # print(res)
        # res2 = []
        # for i in res:
        #     newi = re.search('\t(.*)', i)
        #     res2.append(newi.group().replace('\t', '').replace('"', ''))
        h = open('new_drug.txt', encoding='utf-8')
        res = h.read().splitlines()
        medcine = []
        reason = []
        num = 0
        for i in res:
            # a = re.search('.*?' + '\s?' + i + '\s?' + '.*?\.', txt, re.I | re.IGNORECASE) #版本1，废弃
            a = re.search('.*?' + '\s?' + i + '\s?', self.txt)
            if a:
                print(i)
                searchtxt = a
                print(searchtxt.group())
                medcine.append(i)
                reason.append(searchtxt.group())
            print('\r'+str(num / len(res) * 100) + '%',end='')
            num += 1
        return [medcine,reason]

    # 样本量
    def search_num(self):
        searchtxt = ''
        a = re.search('.*?' + '\s' + '\d+\spatients' + '\s?' + '.*?\D\.', self.txt, re.IGNORECASE)
        if a:
            searchtxt = a
            return searchtxt.group()
        else:
            return 'No Match'

    # 线数
    def search_phase(self):
        a = re.search('.*?' + '\s' + 'phase I' + '\s?' + '.*?\D\.', self.txt,  re.IGNORECASE)
        if a:
            searchtxt = a
            return searchtxt.group()
        else:
            return 'No Match'

    # 实验目的
    def search_purpose(self):
        a = re.search('.*?' + '\s' + 'background' + '\s?' + '.*?\..*?\.', self.txt,  re.I | re.IGNORECASE)
        b = re.search('background' + '\s' + '.*?\..*?\.', self.txt, re.I | re.IGNORECASE)
        if b:
            searchtxt = b
            return searchtxt.group()
        elif a:
            searchtxt = a
            return searchtxt.group()
        else:
            return 'No Match'

    # 研究结果
    def search_res(self):
        a = re.search('.*?' + '\s' + 'findings' + '\s?' + '.*?\..*?\D\.', self.txt, re.I | re.IGNORECASE)
        b = re.search('\s' + 'results' + '\s?' + '.*?\..*?\D\.', self.txt, re.I | re.IGNORECASE)
        if b:
            searchtxt = b
            return searchtxt.group()

        elif a:
            searchtxt = a
            return searchtxt.group()
        else:
            return 'No Match'

    # 研究结论
    def search_con(self):
        a = re.search('conclusions' + '\s?' + '.*?\D\.', self.txt, re.I | re.IGNORECASE)
        b = re.search('.*?' + '\s' + 'INTERPRETATION' + '\s?' + '.*?\D\.', self.txt, re.I | re.IGNORECASE)
        if a:
            searchtxt = a
            return searchtxt.group()

        elif b:
            searchtxt = b
            return searchtxt.group()
        else:
            return 'No Match'

    def search_disscussion(self):
        h = open('discussion.txt', encoding='gbk')
        res = h.read().splitlines()
        # print(res)
        temp = []
        for i in res:
            a = re.search('[)](.*)([(])?', i)
            b = a.group().replace(')', '')
            c = re.sub('（.*?）', '', b)
            temp.append(c)
        keyword = []
        reason = []
        num = 0
        for i in temp:
            a = re.search('.*?' + '\s' + i + '\s?' + '.*?\.', self.txt, re.I | re.IGNORECASE)
            if a:
                searchtxt = a
                keyword.append(i)
                reason.append(searchtxt.group())
            print('\r'+str(num / len(res) * 100) + '%',end='')
            num += 1
        return [keyword,reason]


    def table_chart(self):
        # table = re.findall('.*?' + '\s' + 'table' + '\s?' + '.*?\.', self.txt, re.I | re.IGNORECASE)
        # chart = re.findall('.*?' + '\s' + 'chart' + '\s?' + '.*?\.', self.txt, re.I | re.IGNORECASE)
        # table2 = re.findall('.*?' + '\s' + 'table' + '\s\d' + '.*\n',self.txt,re.I | re.DOTALL)
        # table.extend(chart)
        # table.extend(table2) #版本1弃用
        #table = re.findall('Table' + '\s?\d' + '.*?\n.*?\.', txt, re.DOTALL) #版本2弃用

        # table = re.findall('.*?\(.*?\s?Table' + '\s? ' + '.*?\)', self.txt)
        # if len(table) ==0:
        #     table = re.findall('.*?\(.*?\s?Table' + '\s? ' + '.*?\)', self.txt,re.I) #版本3
        res = []
        res_table = []
        res_figure = []
        table = re.findall('.*?\(Table' + '\s? ' + '.*?\)', self.txt, re.I)
        for i in table:
            try:
                table_num = re.search('\(table.*?([1-9])[\s,\,,\),A,B,C]', i,re.I).group(1)
                # res_table.append((table_num,i))
                res_table.append(i)
            except Exception:
                print(i)
            else:
                pass
        figure = re.findall('.*?\(figure' + '\s? ' + '.*?\)', self.txt, re.I)
        for i in figure:
            try:
                figure_num = re.search('\(figure.*?([1-9])[\s,\,,\),A,B,C]', i,re.I).group(1)
                # res_figure.append((figure_num,i))
                res_figure.append(i)
            except Exception:
                print(i)

            else:
                pass
        # res = [res_table , res_figure]
        res_table.extend(res_figure)
        return res_table

    def table_chart_new(self):
        """
        更改表格陈述匹配规则，更加宽泛
        :return:
        """
        table = re.findall('.*?\(.*?Table' + '\s' + '.*?\)', self.txt)
        table2 = self.table_chart()
        table.extend(table2)
        return table


    def docx_write(self,document,filename,output_dir,image_name):

        # print('搜寻相关药物')
        # medcine_reason = self.search_medicine(txt)
        #
        # document.add_heading(u'研究相关药物',1)
        # document.add_paragraph(str(medcine_reason[0]))

        patients = self.search_patient()
        document.add_heading(u'研究患者', 1)
        document.add_paragraph(patients)

        num = self.search_num()
        document.add_heading(u'样本量', 1)
        document.add_paragraph(num)

        characteristics = self.search_characteristics()
        document.add_heading(u'基线特征', 1)
        document.add_paragraph(characteristics)

        phase = self.search_phase()
        document.add_heading(u'试验设计', 1)
        document.add_paragraph(phase)

        purpose = self.search_purpose()
        document.add_heading(u'研究背景', 1)
        document.add_paragraph(purpose)

        result = self.search_res()
        document.add_heading(u'研究结果', 1)
        document.add_paragraph(result)

        conclution = self.search_con()
        document.add_heading(u'研究结论', 1)
        document.add_paragraph(conclution)

        # print('搜寻讨论')
        # discussion =self.search_disscussion()
        # document.add_heading(u'讨论关键词及讨论内容', 1)
        # document.add_paragraph(str(discussion))

        new_table = self.table_chart_new()
        document.add_heading(u'表格及图片陈述',1)
        for i in new_table:
            document.add_paragraph(u'>>>>>>>>>>>>>>>>>>')
            # document.add_paragraph(i)
            try:
                newi = i.replace('', '')
                document.add_paragraph(newi)

            except Exception as e:
                print(str(e))
                print(i)

        tableandchart = self.table_chart() #table_chart为二位列表，0位为table，1位为figure

        document.add_heading(u'表格及图片', 1)
        path = 'tmp/'+ image_name +'/table'
        path2 = 'tmp/'+ image_name +'/figure'
        files = os.listdir(path)
        files2 = os.listdir(path2)
        f1_num = 1
        f2_num = 1
        document.add_heading(u'表格', 2)
        for i in files:
            document.add_picture(path+'/'+i,width=Inches(6))
            # for j in tableandchart[0]:
            #     if j[0]==str(f1_num):
            #         try:
            #             newi = j[1].replace('', '')
            #             document.add_paragraph(newi)
            #             document.add_paragraph(u'>>>>>>>>>>>>>>>>>>')
            #         except Exception as e:
            #             print(str(e))
            #             print(j[1])
            #         else:
            #             pass
            f1_num += 1
        document.add_heading(u'图片', 2)
        for i in files2:
            document.add_picture(path2 + '/' + i, width=Inches(6))
            # for j in tableandchart[1]:
            #     if j[0]==str(f2_num):
            #         try:
            #             newi = j[1].replace('', '')
            #             document.add_paragraph(newi)
            #             document.add_paragraph(u'>>>>>>>>>>>>>>>>>>')
            #         except Exception as e:
            #             print(str(e))
            #             print(j[1])
            #         else:
            #             pass
            # f2_num += 1
        # for i in tableandchart:
        #     try:
        #         newi = i.replace('', '')
        #         document.add_paragraph(newi)
        #         document.add_paragraph(u'>>>>>>>>>>>>>>>>>>')
        #     except Exception as e:
        #         print(str(e))
        #         print(i)
        #     else:
        #         pass

        document.save(output_dir+'/result_for_'+filename+'.docx')
        # document.save(filename)

if __name__ == '__main__':
    # 遍历文件
    for i in files:
        document = Document()
        filename = i
        document.add_heading(i, 0)
        f = open(path+'/'+i,'r',encoding='utf-8')
        txt = f.read()
        query_ans = Qurey_answer(txt)
        f.close()
        print('\n'+i+'-----------------------------------------')
        query_ans.docx_write(document)



